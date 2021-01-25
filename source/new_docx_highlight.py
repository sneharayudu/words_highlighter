import docx
import re
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from collections import Counter


doc = Document('Blog Role of RPA during and after COVID-19.docx')

list_pattern = [
				"RPA technology",
				"rpa",
				"robotic process automation",
				"RPA tools",
				"what is rpa",
				"rpa automation",
				"rpa software",
				"rpa service",
				"rpa solutions",
				"rpa companies",
				"rpa platform",
				"rpa process",
				"free rpa tools",
				"rpa tools comparison",
				"rpa vendors",
				"rpa systems",
				"about rpa",
				"types of rpa",
				"rpa capabilities",
				"edgeverve rpa",
				"rpa overview",
				"rpa roles",
				"rpa security",
				"role of rpa",
				"roles in rpa",
				"best rpa",
				"rpa business",
				"rpa automate",
				"rpa servicenow",
				"rpa software vendors",
				"rpa managed services",
				"best rpa solutions",
				"top rpa solutions"
				]

list_pattern = [i.lower().strip() for i in list_pattern]

def remove_duplicate(new_dict):
	res_final = new_dict.copy()
	asending = [i[0] for i in list(new_dict.values())]
	y = [item for item, count in Counter(asending).items() if count > 1]
	for item in y:
	    same = {}
	    for k,v in new_dict.items():
	        if v[0] == item:
	            same[k] = v
	    for k,v in new_dict.items():
	        if v == sorted(same.values())[0]:
	            del res_final[k]
	return res_final




for p in doc.paragraphs:
	res = []
	for pattern in list_pattern:
	    if re.search(pattern, p.text.lower().strip(),re.IGNORECASE):
	        res.extend(re.findall(pattern, p.text.lower().strip(),re.IGNORECASE) )
	# print(res)
	# print(p.text)
	if len(res) > 0:
		runs = list(p.runs)
		p.text = ''
		for run in runs:
			p.style.font.name = run.font.name
			p.style.font.size = run.font.size
			flag = False
			link_chk = re.search("http", run.text.strip().lower(),re.IGNORECASE)
			if run.bold:
				flag = True
				newrun = p.add_run(run.text)
				if run.bold:
				    newrun.bold = True
				if run.italic:
				    newrun.italic = True

			if len(res) == 0 or link_chk:
				flag = True
				newrun = p.add_run(run.text)
				if run.bold:
				    newrun.bold = True
				if run.italic:
				    newrun.italic = True

			if flag == False :
				new_dict  = {}
				for match_word in res:
				    find_the_word = re.finditer(match_word,run.text,re.IGNORECASE)

				    for i,match in enumerate(find_the_word):
				        new_dict[match.group() + str(i)] = (match.start(),match.end())
				# print(new_dict)

				res_final = remove_duplicate(new_dict)
				# print(res_final)
				if len(res_final) > 0:
					val = sorted(res_final.values()) 
					start = val[0][0]
					p.add_run(run.text[0:start])
					l = 0 
					while l < len(res_final):
						# print(run.font.size)
						colored = p.add_run(run.text[val[l][0]:val[l][1]])
						colored.font.highlight_color = WD_COLOR_INDEX.YELLOW
						if l  == (len(res_final) - 1):
							p.add_run(run.text[val[l][1]:len(run.text)+1])
						else:
							p.add_run(run.text[val[l][1]:val[l + 1][0]])
						l += 1
                

doc.save('demo.docx')